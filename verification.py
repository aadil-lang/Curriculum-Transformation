from __future__ import annotations

import csv
import json
import os
import tempfile
from dataclasses import asdict, dataclass
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from agent_engine import PreExtractionUnderstanding
from chat_batches import build_chat_batch_request, run_chat_batches_with_factory
from docx import Document

from config import RuntimePaths, get_runtime_paths, get_settings
from pipeline import DataTransformationPipeline
from schemas import get_extraction_payload_model, get_schema_fields


@dataclass(slots=True)
class OfflineVerificationResult:
    temp_root: str
    runtime_paths: dict[str, str]
    process_results: list[dict[str, str]]
    csv_rows: list[dict[str, str]]
    manual_review_entries: list[dict[str, Any]]
    retry_recovery: dict[str, Any]
    chat_batch: dict[str, Any]
    instruction_batch: dict[str, Any]


class StubExtractionEngine:
    def __init__(self, schema_path: str | None = None) -> None:
        self.schema_path = schema_path

    def extract(self, parsed_document: Any, prior_error_log: str | None = None, region_override: Any | None = None) -> Any:
        payload_model = get_extraction_payload_model(self.schema_path)
        payload = payload_model.model_validate(_build_stub_payload(self.schema_path))
        planning = PreExtractionUnderstanding(
            layout_analysis=["Heading followed by labeled key-value paragraphs."],
            row_formation_logic=["Build one row from the visible labeled fields in the document."],
            column_derivations={spec.name: f"Derived from the stubbed value for '{spec.name}'." for spec in get_schema_fields(self.schema_path)},
            representative_row={spec.name: str(payload.model_dump().get(spec.name, "")) for spec in get_schema_fields(self.schema_path)},
            exclusion_rules=["Ignore unrelated layout noise in the verification fixture."],
        )
        return SimpleNamespace(
            payload_rows=[payload],
            planning=planning,
            layout_analysis=planning.layout_analysis,
            anchoring_plan={"record_title": "document heading", "record_owner": "Owner line"},
        )


class StubCritic:
    def validate(self, row: Any, parsed_document: Any) -> Any:
        return SimpleNamespace(tag="VALID", is_valid=True, issues=[], confidence_notes="")


class MissingKeyExtractor:
    def extract(self, parsed_document: Any, prior_error_log: str | None = None, region_override: Any | None = None) -> Any:
        raise RuntimeError("GEMINI_API_KEY is not configured.")


def _build_verification_inputs(runtime_paths: RuntimePaths) -> None:
    success_doc = Document()
    success_doc.add_heading("Offline Verification Doc", level=1)
    success_doc.add_paragraph("Owner: Verification Vendor")
    success_doc.add_paragraph("Date: 2026-07-05")
    success_doc.add_paragraph("Amount: 42.00")
    success_doc.add_paragraph("Offline verification summary line")
    success_doc.save(runtime_paths.input_dir / "verified_smoke.docx")

    (runtime_paths.input_dir / "bad_reference.webloc").write_text(
        "not a valid property list",
        encoding="utf-8",
    )


def run_offline_smoke_verification() -> OfflineVerificationResult:
    settings = get_settings()
    with tempfile.TemporaryDirectory(prefix="data-agent-selftest-") as temp_root_str:
        temp_root = Path(temp_root_str)
        runtime_paths = get_runtime_paths(temp_root)
        _build_verification_inputs(runtime_paths)

        pipeline = DataTransformationPipeline(
            settings=settings,
            runtime_paths=runtime_paths,
            extractor=StubExtractionEngine(),
            critic=StubCritic(),
        )
        results = pipeline.process_pending_documents()

        csv_rows: list[dict[str, str]] = []
        if runtime_paths.final_csv_path.exists():
            with runtime_paths.final_csv_path.open(newline="", encoding="utf-8") as handle:
                csv_rows = list(csv.DictReader(handle))

        manual_review_entries: list[dict[str, Any]] = []
        if runtime_paths.manual_review_path.exists():
            manual_review_entries = json.loads(runtime_paths.manual_review_path.read_text(encoding="utf-8"))

        retry_recovery = _run_retry_recovery_verification(settings)
        chat_batch = _run_chat_batch_verification(settings)
        instruction_batch = _run_instruction_batch_verification(settings)

        return OfflineVerificationResult(
            temp_root=str(temp_root),
            runtime_paths={
                "input_dir": str(runtime_paths.input_dir),
                "output_dir": str(runtime_paths.output_dir),
                "final_csv_path": str(runtime_paths.final_csv_path),
                "manual_review_path": str(runtime_paths.manual_review_path),
            },
            process_results=[
                {
                    "path": str(result.path),
                    "status": result.status,
                    "message": result.message,
                    "stage": result.stage,
                }
                for result in results
            ],
            csv_rows=csv_rows,
            manual_review_entries=manual_review_entries,
            retry_recovery=retry_recovery,
            chat_batch=chat_batch,
            instruction_batch=instruction_batch,
        )


def _run_retry_recovery_verification(settings: Any) -> dict[str, Any]:
    tracked_env_keys = ("GEMINI_API_KEY", "OPENAI_API_KEY", "ANTHROPIC_API_KEY")
    original_env = {key: os.environ.get(key) for key in tracked_env_keys}

    try:
        for key in tracked_env_keys:
            os.environ.pop(key, None)

        with tempfile.TemporaryDirectory(prefix="data-agent-retry-") as temp_root_str:
            temp_root = Path(temp_root_str)
            runtime_paths = get_runtime_paths(temp_root)

            document = Document()
            document.add_heading("Offline Verification Doc", level=1)
            document.add_paragraph("Owner: Verification Vendor")
            document.add_paragraph("Date: 2026-07-05")
            document.add_paragraph("Amount: 42.00")
            document.add_paragraph("Offline verification summary line")
            document.save(runtime_paths.input_dir / "retry_after_config.docx")

            first_pipeline = DataTransformationPipeline(
                settings=settings,
                runtime_paths=runtime_paths,
                extractor=MissingKeyExtractor(),
                critic=StubCritic(),
            )
            first_results = first_pipeline.process_pending_documents()
            first_state = json.loads(runtime_paths.state_path.read_text(encoding="utf-8"))

            os.environ["GEMINI_API_KEY"] = "dummy-gemini-key"
            os.environ["OPENAI_API_KEY"] = "dummy-openai-key"

            second_pipeline = DataTransformationPipeline(
                settings=settings,
                runtime_paths=runtime_paths,
                extractor=StubExtractionEngine(str(settings.schema_config_path)),
                critic=StubCritic(),
            )
            second_results = second_pipeline.process_pending_documents()
            second_state = json.loads(runtime_paths.state_path.read_text(encoding="utf-8"))

            csv_rows: list[dict[str, str]] = []
            if runtime_paths.final_csv_path.exists():
                with runtime_paths.final_csv_path.open(newline="", encoding="utf-8") as handle:
                    csv_rows = list(csv.DictReader(handle))

            return {
                "first_results": [
                    {
                        "path": str(result.path),
                        "status": result.status,
                        "message": result.message,
                        "stage": result.stage,
                    }
                    for result in first_results
                ],
                "first_state": first_state,
                "second_results": [
                    {
                        "path": str(result.path),
                        "status": result.status,
                        "message": result.message,
                        "stage": result.stage,
                    }
                    for result in second_results
                ],
                "second_state": second_state,
                "csv_rows": csv_rows,
            }
    finally:
        for key, value in original_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value


def _run_chat_batch_verification(settings: Any) -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="data-agent-chat-batch-") as temp_root_str:
        temp_root = Path(temp_root_str)

        sample_csv_path = temp_root / "invoice_template.csv"
        sample_csv_path.write_text("Invoice Number,Vendor,Amount\nINV-001,ACME,100.50\n", encoding="utf-8")

        document_path = temp_root / "invoice_1.docx"
        document = Document()
        document.add_heading("Invoice", level=1)
        document.add_paragraph("Invoice Number: INV-001")
        document.add_paragraph("Vendor: ACME")
        document.add_paragraph("Amount: 100.50")
        document.save(document_path)

        request = build_chat_batch_request(
            name="invoice_batch",
            files=[str(document_path)],
            sample_csv=str(sample_csv_path),
            instructions=None,
            infer_schema=False,
            draft_only=False,
            output_csv_name="invoice_batch_output.csv",
        )

        def pipeline_factory(batch_settings: Any, runtime_paths: RuntimePaths) -> DataTransformationPipeline:
            return DataTransformationPipeline(
                settings=batch_settings,
                runtime_paths=runtime_paths,
                extractor=StubExtractionEngine(str(batch_settings.schema_config_path)),
                critic=StubCritic(),
            )

        results = run_chat_batches_with_factory(request, settings, pipeline_factory=pipeline_factory)
        result = results[0]

        output_csv_path = Path(result.output_csv_path)
        csv_rows: list[dict[str, str]] = []
        if output_csv_path.exists():
            with output_csv_path.open(newline="", encoding="utf-8") as handle:
                csv_rows = list(csv.DictReader(handle))

        schema_payload = json.loads(Path(result.schema_path).read_text(encoding="utf-8"))

        return {
            "results": [asdict(result)],
            "schema": schema_payload,
            "csv_rows": csv_rows,
        }


def _run_instruction_batch_verification(settings: Any) -> dict[str, Any]:
    request = build_chat_batch_request(
        name="instruction_only_batch",
        files=[],
        sample_csv=None,
        instructions="Create columns for Claim ID, Claim Date, Customer Name, Amount, and Claim Summary.",
        infer_schema=True,
        draft_only=True,
        output_csv_name=None,
    )

    results = run_chat_batches_with_factory(request, settings)
    result = results[0]

    schema_payload = json.loads(Path(result.schema_path).read_text(encoding="utf-8"))
    template_rows: list[list[str]] = []
    sample_template_path = Path(result.schema_path).with_name("sample_output_template.csv")
    if sample_template_path.exists():
        with sample_template_path.open(newline="", encoding="utf-8") as handle:
            template_rows = list(csv.reader(handle))

    return {
        "results": [asdict(result)],
        "schema": schema_payload,
        "sample_template_rows": template_rows,
    }


def _build_stub_payload(schema_path: str | None) -> dict[str, Any]:
    payload: dict[str, Any] = {}
    for spec in get_schema_fields(schema_path):
        if spec.field_type == "number":
            value: Any = 42.0
            citation = "42.00"
        elif spec.field_type == "integer":
            value = 7
            citation = "7"
        elif spec.field_type == "boolean":
            value = True
            citation = "true"
        else:
            value = f"sample_{spec.name}"
            citation = value
        payload[spec.name] = value
        payload[f"{spec.name}_source_citation"] = citation
    return payload
