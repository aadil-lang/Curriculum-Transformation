from __future__ import annotations

import csv
import json
import os
import tempfile
from dataclasses import asdict, dataclass
from pathlib import Path
from types import SimpleNamespace
from typing import Any

from extractor import PreExtractionUnderstanding
from batch_runner import build_chat_batch_request, run_batches_with_factory
from docx import Document

from config import RuntimePaths, get_runtime_paths, get_settings
from pipeline import DataTransformationPipeline
from schemas import get_extraction_payload_model, get_schema_fields


@dataclass(slots=True)
class OfflineVerificationResult:
    temp_root: str
    runtime_paths: dict[str, str]
    process_results: list[dict[str, str]]
    csv_rows: list[dict[str, str]]
    manual_review_entries: list[dict[str, Any]]
    chat_batch: dict[str, Any]
    instruction_batch: dict[str, Any]


class StubExtractionEngine:
    def __init__(self, schema_path: str | None = None) -> None:
        self.schema_path = schema_path

    def extract(self, parsed_document: Any, prior_error_log: str | None = None, region_override: Any | None = None) -> Any:
        from config import get_settings

        payload_model = get_extraction_payload_model(
            self.schema_path, include_citations=get_settings().extraction_citations_enabled
        )
        payload = payload_model.model_validate(_build_stub_payload(self.schema_path))
        planning = PreExtractionUnderstanding(
            layout_analysis=["Heading followed by labeled key-value paragraphs."],
            row_formation_logic=["Build one row from the visible labeled fields in the document."],
            column_derivations={spec.name: f"Derived from the stubbed value for '{spec.name}'." for spec in get_schema_fields(self.schema_path)},
            representative_row={spec.name: str(payload.model_dump().get(spec.name, "")) for spec in get_schema_fields(self.schema_path)},
            exclusion_rules=["Ignore unrelated layout noise in the verification fixture."],
        )
        return SimpleNamespace(
            payload_rows=[payload],
            planning=planning,
            layout_analysis=planning.layout_analysis,
            anchoring_plan={"record_title": "document heading", "record_owner": "Owner line"},
        )


class StubCritic:
    def validate(self, row: Any, parsed_document: Any) -> Any:
        return SimpleNamespace(tag="VALID", is_valid=True, issues=[], confidence_notes="")


def _add_curriculum_fixture_content(document: "Document") -> None:
    """Curriculum-shaped fixture content matching the stub payload values.

    offline-verify runs the REAL reviewer, which rejects values not supported by the
    source. So the fixture must contain the curriculum content the stub 'extracts'
    (Grade 5 Science, Forces and Motion, standard SC.5.P.13.1).
    """
    document.add_heading("Grade 5 Science Standards", level=1)
    document.add_paragraph("Subject: Science")
    document.add_paragraph("Grade Level: Elementary School (Grade 5)")
    document.add_paragraph("Domain: Physical Science")
    document.add_paragraph("Topic: Forces and Motion")
    document.add_paragraph(
        "SC.5.P.13.1 Identify familiar forces that cause objects to move, such as pushes or pulls."
    )


def _build_verification_inputs(runtime_paths: RuntimePaths) -> None:
    success_doc = Document()
    _add_curriculum_fixture_content(success_doc)
    success_doc.save(runtime_paths.input_dir / "verified_smoke.docx")

    (runtime_paths.input_dir / "bad_reference.webloc").write_text(
        "not a valid property list",
        encoding="utf-8",
    )


def run_offline_smoke_verification() -> OfflineVerificationResult:
    settings = get_settings()
    with tempfile.TemporaryDirectory(prefix="data-agent-selftest-") as temp_root_str:
        temp_root = Path(temp_root_str)
        runtime_paths = get_runtime_paths(temp_root)
        _build_verification_inputs(runtime_paths)

        pipeline = DataTransformationPipeline(
            settings=settings,
            runtime_paths=runtime_paths,
            extractor=StubExtractionEngine(),
            critic=StubCritic(),
        )
        results = pipeline.process_pending_documents()

        csv_rows: list[dict[str, str]] = []
        if runtime_paths.final_csv_path.exists():
            with runtime_paths.final_csv_path.open(newline="", encoding="utf-8") as handle:
                csv_rows = list(csv.DictReader(handle))

        manual_review_entries: list[dict[str, Any]] = []
        if runtime_paths.manual_review_path.exists():
            manual_review_entries = json.loads(runtime_paths.manual_review_path.read_text(encoding="utf-8"))

        chat_batch = _run_chat_batch_verification(settings)
        instruction_batch = _run_instruction_batch_verification(settings)

        return OfflineVerificationResult(
            temp_root=str(temp_root),
            runtime_paths={
                "input_dir": str(runtime_paths.input_dir),
                "output_dir": str(runtime_paths.output_dir),
                "final_csv_path": str(runtime_paths.final_csv_path),
                "manual_review_path": str(runtime_paths.manual_review_path),
            },
            process_results=[
                {
                    "path": str(result.path),
                    "status": result.status,
                    "message": result.message,
                    "stage": result.stage,
                }
                for result in results
            ],
            csv_rows=csv_rows,
            manual_review_entries=manual_review_entries,
            chat_batch=chat_batch,
            instruction_batch=instruction_batch,
        )


def _run_chat_batch_verification(settings: Any) -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="data-agent-chat-batch-") as temp_root_str:
        temp_root = Path(temp_root_str)

        sample_csv_path = temp_root / "invoice_template.csv"
        sample_csv_path.write_text("Invoice Number,Vendor,Amount\nINV-001,ACME,100.50\n", encoding="utf-8")

        document_path = temp_root / "invoice_1.docx"
        document = Document()
        document.add_heading("Invoice", level=1)
        document.add_paragraph("Invoice Number: INV-001")
        document.add_paragraph("Vendor: ACME")
        document.add_paragraph("Amount: 100.50")
        document.save(document_path)

        request = build_chat_batch_request(
            name="invoice_batch",
            files=[str(document_path)],
            sample_csv=str(sample_csv_path),
            instructions=None,
            infer_schema=False,
            draft_only=False,
            output_csv_name="invoice_batch_output.csv",
        )

        def pipeline_factory(batch_settings: Any, runtime_paths: RuntimePaths) -> DataTransformationPipeline:
            return DataTransformationPipeline(
                settings=batch_settings,
                runtime_paths=runtime_paths,
                extractor=StubExtractionEngine(str(batch_settings.schema_config_path)),
                critic=StubCritic(),
            )

        results = run_batches_with_factory(request, settings, pipeline_factory=pipeline_factory)
        result = results[0]

        output_csv_path = Path(result.output_csv_path)
        csv_rows: list[dict[str, str]] = []
        if output_csv_path.exists():
            with output_csv_path.open(newline="", encoding="utf-8") as handle:
                csv_rows = list(csv.DictReader(handle))

        schema_payload = json.loads(Path(result.schema_path).read_text(encoding="utf-8"))

        return {
            "results": [asdict(result)],
            "schema": schema_payload,
            "csv_rows": csv_rows,
        }


def _run_instruction_batch_verification(settings: Any) -> dict[str, Any]:
    request = build_chat_batch_request(
        name="instruction_only_batch",
        files=[],
        sample_csv=None,
        instructions="Create columns for Claim ID, Claim Date, Customer Name, Amount, and Claim Summary.",
        infer_schema=True,
        draft_only=True,
        output_csv_name=None,
    )

    results = run_batches_with_factory(request, settings)
    result = results[0]

    schema_payload = json.loads(Path(result.schema_path).read_text(encoding="utf-8"))
    template_rows: list[list[str]] = []
    sample_template_path = Path(result.schema_path).with_name("sample_output_template.csv")
    if sample_template_path.exists():
        with sample_template_path.open(newline="", encoding="utf-8") as handle:
            template_rows = list(csv.reader(handle))

    return {
        "results": [asdict(result)],
        "schema": schema_payload,
        "sample_template_rows": template_rows,
    }


def _build_stub_payload(schema_path: str | None) -> dict[str, Any]:
    # Only emit `<field>_source_citation` keys when the actual payload model has them
    # (i.e. citations enabled). The model is extra="forbid", so injecting citation keys
    # while citations are disabled fails validation — the stub must match the real model.
    from config import get_settings

    include_citations = get_settings().extraction_citations_enabled
    payload_model = get_extraction_payload_model(schema_path, include_citations=include_citations)
    model_fields = set(payload_model.model_fields.keys())

    # Contract-valid curriculum values that MATCH the verification fixture document
    # (see _build_verification_inputs). The real reviewer runs during offline-verify and
    # rejects placeholder junk, so these must look like genuine extracted curriculum rows
    # — e.g. grade_level must be one of the allowed buckets.
    curriculum_values: dict[str, str] = {
        "source": "https://example.edu/curriculum/science/grade-5",
        "grade_level": "Elementary School",
        "display_grade": "5",
        "grade_number": "5",
        "subject": "Science",
        "domain": "Physical Science",
        "topic": "Forces and Motion",
        "standard_code": "SC.5.P.13.1",
        "display_standard_code": "SC.5.P.13.1",
        "description": "Identify familiar forces that cause objects to move, such as pushes or pulls.",
    }

    payload: dict[str, Any] = {}
    for spec in get_schema_fields(schema_path):
        if spec.name in curriculum_values:
            value: Any = curriculum_values[spec.name]
            citation = str(value)
        elif spec.field_type == "number":
            value = 42.0
            citation = "42.00"
        elif spec.field_type == "integer":
            value = 7
            citation = "7"
        elif spec.field_type == "boolean":
            value = True
            citation = "true"
        else:
            # Optional/unknown fields (l3/l4/l5, czi_standard_code) stay blank — valid and
            # matches the sample-blank contract rather than inventing hierarchy levels.
            value = ""
            citation = ""
        payload[spec.name] = value
        citation_key = f"{spec.name}_source_citation"
        if citation_key in model_fields:
            payload[citation_key] = citation
    return payload
